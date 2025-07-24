from docx import Document
from docx.shared import Pt
import os

# Create a new document
doc = Document()

# Utility functions
def add_centered_heading(text, size=20):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    paragraph.alignment = 1  # Center alignment

def add_bold_heading(text):
    doc.add_heading(text, level=2)

def add_bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def add_normal_paragraph(text):
    doc.add_paragraph(text)

# Resume Content
add_centered_heading("KEMIDI MANISH")

add_normal_paragraph("📞 7036969854 | ✉️ kemidimanish04@gmail.com")
add_normal_paragraph("🌐 LinkedIn: https://www.linkedin.com/in/kemidimanish04")
add_normal_paragraph("💻 GitHub: https://github.com/manish2804-hub")
add_normal_paragraph("🏠 Address: Boduppal, Hyderabad")

add_bold_heading("Profile")
add_normal_paragraph(
    "Motivated and detail-oriented Computer Science undergraduate with strong interests in Python, Web Development, and Cloud Computing. "
    "Currently focused on enhancing my technical and problem-solving abilities by building real-world projects and contributing to open-source. "
    "Eager to work in collaborative environments and drive innovative software solutions.\n\n"
    "I’m looking forward to working with amazing teams and building impactful solutions.\n\n"
    "Regards & Thank you,\nKemidi Manish"
)

add_bold_heading("Area of Interests")
add_bullets([
    "Web Development (Frontend & Backend)",
    "Cloud Computing (Basics of AWS, Azure)",
    "Machine Learning & Data Science",
    "Open Source Contributions",
    "Git & Version Control",
    "WordPress Development"
])

add_bold_heading("Skills")
add_bullets([
    "Languages: Python, C (Basics), Java (Basics)",
    "Web Technologies: HTML, CSS, JavaScript, PHP",
    "Databases: SQL, MySQL",
    "Tools: Git, GitHub, TensorFlow, Scikit-learn",
    "Platforms: Windows OS, WordPress (Beginner)"
])

add_bold_heading("Education Qualification")
add_normal_paragraph(
    "• B.Tech in CSE (2022–2026) – Kommuri Pratap Reddy Institute of Technology – CGPA: 8.0 (till 2nd year)\n"
    "• Intermediate – MPC (2020–2022) – Narayana Junior College – GPA: 8.7/10\n"
    "• SSC – SR Digi School – GPA: 9.5/10"
)

add_bold_heading("Projects")
add_normal_paragraph(
    "Anomaly Detection in IoT Sensor Data using ML\n"
    "- Used Python, TensorFlow, Scikit-learn, NumPy, Pandas\n"
    "- Built a system to detect unusual patterns in sensor readings"
)

add_bold_heading("Certifications")
add_bullets([
    "Microsoft Copilot – Certified via LinkedIn",
    "Generative AI Course – Completed",
    "Machine Learning with IoT – SAK Informatics"
])

add_bold_heading("Achievements & Strengths")
add_bullets([
    "Active participant in technical & non-technical college events",
    "Strong communication and collaboration skills",
    "Fast learner with a passion for problem-solving"
])

add_bold_heading("Declaration")
add_normal_paragraph(
    "I hereby declare that the information provided above is true to the best of my knowledge and belief.\n\n"
    "Signature: Kemidi Manish\nDate:\nPlace: Hyderabad"
)

# Save the resume
filename = "Kemidi_Manish_Resume.docx"
doc.save(filename)
print(f"Resume saved as '{filename}'")

# Open the resume (Windows only)
os.startfile(filename)
